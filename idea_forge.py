import streamlit as st
import google.generativeai as genai
from dotenv import load_dotenv
import os

# Word document generation
try:
    from docx import Document
    from io import BytesIO
    WORD_EXPORT_AVAILABLE = True
except ImportError:
    WORD_EXPORT_AVAILABLE = False
    st.warning("python-docx is not installed. Word export functionality will be disabled.")

# Load environment variables from .env
load_dotenv(".env")

# Configure Gemini API
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))

# Define Tamil Nadu districts of focus

TN_DISTRICTS = [
    "Ariyalur", "Chengalpattu", "Chennai", "Coimbatore", "Cuddalore", "Dharmapuri", "Dindigul",
    "Erode", "Kallakurichi", "Kanchipuram", "Kanyakumari", "Karur", "Krishnagiri", "Madurai",
    "Nagapattinam", "Namakkal", "Nilgiris", "Perambalur", "Pudukkottai", "Ramanathapuram",
    "Ranipet", "Salem", "Sivaganga", "Tenkasi", "Thanjavur", "Theni", "Thoothukudi", "Tiruchirappalli",
    "Tirunelveli", "Tirupathur", "Tiruppur", "Tiruvallur", "Tiruvannamalai", "Tiruvarur",
    "Vellore", "Viluppuram", "Virudhunagar"
]

# Placeholder for district information
def get_district_info(district):
    district_info = {
        "Ariyalur": "Known for its cement industries and rich limestone resources",
        "Chengalpattu": "Emerging IT hub and automobile manufacturing",
        "Chennai": "Capital city known for IT, automobile, and education sectors",
        "Coimbatore": "Hub for textiles, engineering, and manufacturing",
        "Cuddalore": "Industrial town with chemical and port-based industries",
        "Dharmapuri": "Famous for horticulture and mango cultivation",
        "Dindigul": "Known for lock manufacturing and agriculture",
        "Erode": "Known for turmeric production and textile industry",
        "Kallakurichi": "Agriculture and sugar mills are prominent",
        "Kanchipuram": "Famous for silk weaving and temples",
        "Kanyakumari": "Tourism and wind energy production",
        "Karur": "Known for textile exports and bus body manufacturing",
        "Krishnagiri": "Major producer of mangoes and granite",
        "Madurai": "Cultural hub and known for jasmine cultivation",
        "Nagapattinam": "Fisheries and port-based activities",
        "Namakkal": "Known for poultry farming and transport services",
        "Nilgiris": "Famous for tea plantations and tourism",
        "Perambalur": "Cement industries and agricultural activities",
        "Pudukkottai": "Agriculture and small-scale industries",
        "Ramanathapuram": "Fisheries and seaweed cultivation",
        "Ranipet": "Leather tanning and manufacturing hub",
        "Salem": "Known for steel production and mango cultivation",
        "Sivaganga": "Famous for agriculture and temple tourism",
        "Tenkasi": "Tourism and water-based resources",
        "Thanjavur": "Rice bowl of Tamil Nadu and cultural heritage",
        "Theni": "Known for agriculture and cardamom cultivation",
        "Thoothukudi": "Major port and salt production hub",
        "Tiruchirappalli": "Industrial hub and educational institutions",
        "Tirunelveli": "Known for Halwa, wind energy, and temples",
        "Tirupathur": "Leather industry and small-scale enterprises",
        "Tiruppur": "Knitwear capital of India, focused on garment exports",
        "Tiruvallur": "Automobile and heavy industries",
        "Tiruvannamalai": "Pilgrimage center and agricultural activities",
        "Tiruvarur": "Agriculture, especially paddy cultivation",
        "Vellore": "Renowned for leather industries and educational institutions",
        "Viluppuram": "Agriculture and cashew processing",
        "Virudhunagar": "Famous for fireworks and matchbox industries"
    }
    return district_info.get(district, "Information not available")

# Generate multiple business ideas based on user inputs
def generate_business_ideas(user_data):
    prompt = f"""
    The user is from {user_data['district']}, a region known for {get_district_info(user_data['district'])}. 
    They have the following profile:
    - Qualifications: {user_data['qualifications']}
    - Budget: {user_data['budget']} INR
    - Interests: {user_data['interests']}
    
    Please generate 3 business ideas that are highly likely to succeed in {user_data['district']}. Each idea should:
    1. Align with the user's qualifications and interests
    2. Be feasible within the given budget
    3. Leverage opportunities unique to {user_data['district']}
    4. Have a high probability of success based on local market conditions
    
    For each business idea, provide:
    1. Business Title
    2. Brief Description
    3. Required Resources
    4. Potential Challenges
    5. Marketing Strategy
    6. Estimated Timeline
    7. Growth Potential
    8. Actionable Steps to Start
    9. Suggested Budget Distribution
    """
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(prompt)
    return response.text

# Chatbot function to assist user with project development
def chatbot(user_input, user_data):
    prompt = f"""
    You are a helpful assistant specializing in project development and entrepreneurship in Tamil Nadu, India. 
    The user is from {user_data['district']}, which is known for: {get_district_info(user_data['district'])}
    
    User's profile:
    - Qualifications: {user_data['qualifications']}
    - Budget: {user_data['budget']} INR
    - Interests: {user_data['interests']}
    
    Please respond to the following question or request, taking into account the user's background and local context: 
    {user_input}
    
    Provide practical, actionable advice that is relevant to the user's situation and the local market conditions.
    """
    
    model = genai.GenerativeModel('gemini-pro')
    response = model.generate_content(prompt)
    return response.text

# Function to calculate budget distribution
def calculate_budget_distribution(budget):
    # Example allocation: adjust as needed for specific use cases
    rent = budget * 0.30  # 30% for rent/space
    equipment = budget * 0.25  # 25% for equipment and tools
    marketing = budget * 0.15  # 15% for marketing and promotions
    salaries = budget * 0.20  # 20% for initial staff salaries
    misc = budget * 0.10  # 10% for miscellaneous/other expenses
    return {
        "Rent/Space": rent,
        "Equipment/Tools": equipment,
        "Marketing": marketing,
        "Salaries": salaries,
        "Miscellaneous": misc
    }

# Function to create Word document
def create_word_document(user_data, business_ideas, budget_distribution):
    if not WORD_EXPORT_AVAILABLE:
        st.error("Word export is not available. Please install python-docx to use this feature.")
        return None

    doc = Document()
    doc.add_heading(f"Business Ideas for {user_data['name']}", 0)

    # Add user information
    doc.add_heading("User Information", level=1)
    doc.add_paragraph(f"Name: {user_data['name']}")
    doc.add_paragraph(f"District: {user_data['district']}")
    doc.add_paragraph(f"Qualifications: {user_data['qualifications']}")
    doc.add_paragraph(f"Budget: ₹{user_data['budget']:,}")
    doc.add_paragraph(f"Interests: {user_data['interests']}")

    # Add business ideas
    doc.add_heading("Generated Business Ideas", level=1)
    doc.add_paragraph(business_ideas)

    # Add budget distribution
    doc.add_heading("Budget Distribution", level=1)
    for category, amount in budget_distribution.items():
        doc.add_paragraph(f"{category}: ₹{amount:,.2f}")

    # Save the document to a BytesIO object
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Streamlit UI
st.title("Business Ideas Generator for Tamil Nadu")

# User Input Section
st.header("Your Information")
name = st.text_input("Your name:")
district = st.selectbox("Select your district:", TN_DISTRICTS)
qualifications = st.text_input("Your qualifications:")
budget = st.number_input("Your budget (in INR):", min_value=0, value=100000, step=10000)
interests = st.text_area("Your personal interests:")

# Store user data for business idea generation
user_data = {
    "name": name,
    "district": district,
    "qualifications": qualifications,
    "budget": budget,
    "interests": interests
}

# Business Idea Generation Section
st.header("Generate Business Ideas")

if st.button("Generate Business Ideas"):
    with st.spinner("Generating business ideas..."):
        business_ideas = generate_business_ideas(user_data)
    st.write(business_ideas)

    # Show budget distribution for each business idea
    st.header("Budget Distribution")
    budget_distribution = calculate_budget_distribution(budget)
    
    st.write("Here's how your budget could be distributed:")
    for category, amount in budget_distribution.items():
        st.write(f"{category}: ₹{amount:,.2f}")

    # Generate Word document
    if WORD_EXPORT_AVAILABLE:
        doc_io = create_word_document(user_data, business_ideas, budget_distribution)
        
        # Provide download button for the Word document
        if doc_io:
            st.download_button(
                label="Download Business Ideas as Word Document",
                data=doc_io.getvalue(),
                file_name=f"business_ideas_for_{name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    else:
        st.warning("Word export is not available. Install python-docx to enable this feature.")

# Chatbot Section for Additional Guidance
st.header("Business Development Assistant")
user_input = st.text_input("Ask a question about your business idea or entrepreneurship in your area:")

if user_input:
    with st.spinner("Generating response..."):
        response = chatbot(user_input, user_data)
    st.write("Assistant:", response)
